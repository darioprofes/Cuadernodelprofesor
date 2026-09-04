# ==========================================================
# Anonimizador de documentos -- copia de escritorio
# ==========================================================
#
# Copia manual de api/app/services/anonimizador.py (backend web), con dos
# adaptaciones necesarias por correr aquí congelado con PyInstaller en vez
# de interpretado dentro de un contenedor -- mantener sincronizadas a mano,
# mismo criterio que horario_pdf.py/educastur_client.py en este mismo
# sidecar (ver README.md):
#
# 1. `import es_core_news_md; nlp = es_core_news_md.load()` en vez de
#    `spacy.load("es_core_news_md")` -- la carga por nombre depende de que
#    importlib.metadata pueda resolver el paquete instalado por su nombre
#    de distribución, algo que no funciona dentro del .exe congelado
#    (confirmado en real: revienta con "Can't find model 'es_core_news_md'"
#    nada más arrancar). Importar el paquete del modelo directamente y
#    llamar a su propio .load() es el patrón que la propia documentación de
#    spaCy recomienda para apps congeladas (PyInstaller/cx_Freeze) y
#    funciona sin cambios en el resto del módulo.
# 2. _BASE ya no es la carpeta del propio fichero fuente (que en PyInstaller
#    --onedir vive empaquetado dentro de PYZ-00.pyz, no suelto en disco) --
#    apunta a sys._MEIPASS, que en modo --onedir es la propia carpeta
#    dist/python-helper/ donde main.py deja materias_oficiales.json y
#    neae_terminos.json como datos (ver --add-data en el build). Con
#    sys._MEIPASS ausente (ejecutando este módulo suelto, sin congelar,
#    para probarlo) cae de vuelta a la carpeta del fichero, igual que hacía
#    el original.
#
# El resto del módulo (detección de candidatos, generación de códigos,
# anonimizar/anonimizar_docx/reintegrar_docx) es idéntico al original.

import json
import re
import secrets
import sys
from pathlib import Path

import es_core_news_md

_BASE = Path(getattr(sys, "_MEIPASS", None) or Path(__file__).parent)
_nlp = es_core_news_md.load()

_MATERIAS_CONOCIDAS = {
    m.lower() for m in json.loads((_BASE / "materias_oficiales.json").read_text(encoding="utf-8"))
}

_TERMINOS_NEAE_CONOCIDOS = {
    t.lower() for t in json.loads((_BASE / "neae_terminos.json").read_text(encoding="utf-8"))
}

_PATRON_DNI = re.compile(r"\b\d{8}[A-Za-z]\b")
_PATRON_DIRECCION = re.compile(
    r"\b(?:[Cc]alle|[Aa]venida|[Aa]vda\.?|[Pp]laza|[Pp]aseo|[Cc]/)\s+[^\n,]+?\s+\d+"
    r"(?:,\s*\d+º?[A-Za-z]?)?"
)
_PATRON_CODIGO_POSTAL = re.compile(r"\b\d{5}\b")
_PATRON_CENTRO = re.compile(r"\b(?:IES|CEIP|CRA|CPI|EOI|CIFP|Colegio|Instituto)\b[^\n,\.]{0,30}")
_PATRON_NIVEL = re.compile(
    r"\b[1-4]º\s*(?:de\s+)?ESO(?:\s+[A-ZÑ](?!\w))?"
    r"|\b[1-2]º\s*(?:de\s+)?Bachillerato(?:\s+[A-ZÑ](?!\w))?"
)
_CARGOS_UNICOS = [
    "jefe de estudios", "jefa de estudios", "director", "directora",
    "secretario", "secretaria", "orientador", "orientadora",
]

INSTRUCCION_IA_ONLINE = (
    "IMPORTANTE: este documento contiene códigos con el formato PERS_XXXXXX y "
    "GRUPO_XXXXXX (por ejemplo: PERS_A63425, GRUPO_77DF7E). Son sustitutos de "
    "datos personales reales, generados para proteger la privacidad. NO los "
    "traduzcas, no los modifiques, no los interpretes ni les des otro "
    "significado -- consérvalos EXACTAMENTE tal cual aparecen, letra por "
    "letra, en tu respuesta. Los necesito intactos para reconstruir el "
    "documento con los datos reales después."
)


def _detectar_candidatos(texto):

    doc = _nlp(texto)

    candidatos = [(ent.start_char, ent.end_char, ent.text)
                  for ent in doc.ents if ent.label_ in ("PER", "ORG", "LOC")]

    for patron in (_PATRON_DNI, _PATRON_DIRECCION, _PATRON_CODIGO_POSTAL):
        for m in patron.finditer(texto):
            candidatos.append((m.start(), m.end(), m.group()))

    for m in _PATRON_CENTRO.finditer(texto):
        recortado = m.group().rstrip()
        candidatos.append((m.start(), m.start() + len(recortado), recortado))

    niveles_detectados = [(m.start(), m.end(), m.group()) for m in _PATRON_NIVEL.finditer(texto)]
    candidatos += niveles_detectados

    for cargo in _CARGOS_UNICOS:
        for m in re.finditer(r"\b" + re.escape(cargo) + r"\b", texto, re.IGNORECASE):
            candidatos.append((m.start(), m.end(), m.group()))

    return candidatos, niveles_detectados


def _recortar_saltos_de_linea(candidatos):

    recortados = []

    for inicio, fin, texto in candidatos:
        if "\n" in texto:
            texto = texto.split("\n")[0].rstrip()
            fin = inicio + len(texto)
            if not texto.strip():
                continue
        recortados.append((inicio, fin, texto))

    return recortados


def _resolver_solapamientos(candidatos):

    ordenados = sorted(candidatos, key=lambda c: (c[0], -(c[1] - c[0])))

    resueltos = []
    fin_ocupado_hasta = -1

    for inicio, fin, texto in ordenados:
        if inicio < fin_ocupado_hasta:
            continue
        resueltos.append((inicio, fin, texto))
        fin_ocupado_hasta = fin

    return resueltos


def _generar_codigos(texto):

    candidatos, niveles_detectados = _detectar_candidatos(texto)
    candidatos = _recortar_saltos_de_linea(candidatos)

    candidatos = [
        (i, f, t) for i, f, t in candidatos
        if t.strip().lower() not in _MATERIAS_CONOCIDAS
        and t.strip().lower() not in _TERMINOS_NEAE_CONOCIDOS
    ]

    candidatos = _resolver_solapamientos(candidatos)

    textos_nivel = {texto_n for _, _, texto_n in niveles_detectados}

    codigo_de = {}
    mapa_real = {}
    for _, _, texto_candidato in candidatos:
        if texto_candidato not in codigo_de:
            prefijo = "GRUPO" if texto_candidato in textos_nivel else "PERS"
            codigo = "%s_%s" % (prefijo, secrets.token_hex(3).upper())
            codigo_de[texto_candidato] = codigo
            mapa_real[codigo] = texto_candidato

    return candidatos, codigo_de, mapa_real


def anonimizar(texto):

    candidatos, codigo_de, mapa_real = _generar_codigos(texto)

    anonimizado = texto
    for inicio, fin, texto_candidato in sorted(candidatos, key=lambda c: -c[0]):
        anonimizado = anonimizado[:inicio] + codigo_de[texto_candidato] + anonimizado[fin:]

    for texto_candidato, codigo in sorted(codigo_de.items(), key=lambda x: -len(x[0])):
        anonimizado = anonimizado.replace(texto_candidato, codigo)

    documento_para_pegar = INSTRUCCION_IA_ONLINE + "\n\n" + anonimizado

    return documento_para_pegar, mapa_real


_PATRON_CODIGO_GENERICO = re.compile(r"\b(?:PERS|GRUPO)_[0-9A-F]{6}\b")


def _sustituir_codigos_en_parrafo(paragraph, patron, mapa):

    for run in paragraph.runs:
        if run.text:
            run.text = patron.sub(lambda m: mapa[m.group(0)], run.text)


def _sustituir_textos_en_parrafo(paragraph, pares_texto_codigo):

    for run in paragraph.runs:
        if not run.text:
            continue
        for texto_candidato, codigo in pares_texto_codigo:
            if texto_candidato in run.text:
                run.text = run.text.replace(texto_candidato, codigo)


def anonimizar_docx(contenido_bytes):

    import io

    from docx import Document

    documento = Document(io.BytesIO(contenido_bytes))

    bloques_texto = [p.text for p in documento.paragraphs]
    bloques_texto += [
        celda.text
        for tabla in documento.tables
        for fila in tabla.rows
        for celda in fila.cells
    ]
    texto_completo = "\n".join(bloques_texto)

    _, codigo_de, mapa_real = _generar_codigos(texto_completo)

    if codigo_de:
        pares = sorted(codigo_de.items(), key=lambda x: -len(x[0]))

        for parrafo in documento.paragraphs:
            _sustituir_textos_en_parrafo(parrafo, pares)

        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        _sustituir_textos_en_parrafo(parrafo, pares)

    if documento.paragraphs:
        documento.paragraphs[0].insert_paragraph_before(INSTRUCCION_IA_ONLINE)
    else:
        documento.add_paragraph(INSTRUCCION_IA_ONLINE)

    buffer = io.BytesIO()
    documento.save(buffer)

    return buffer.getvalue(), mapa_real


def reintegrar_docx(contenido_bytes, mapa):

    import io

    from docx import Document

    documento = Document(io.BytesIO(contenido_bytes))

    if mapa:
        patron = re.compile("|".join(re.escape(codigo) for codigo in mapa))

        for parrafo in documento.paragraphs:
            _sustituir_codigos_en_parrafo(parrafo, patron, mapa)

        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        _sustituir_codigos_en_parrafo(parrafo, patron, mapa)

    buffer = io.BytesIO()
    documento.save(buffer)

    texto_completo = "\n".join(p.text for p in documento.paragraphs)
    texto_completo += "\n".join(
        celda.text
        for tabla in documento.tables
        for fila in tabla.rows
        for celda in fila.cells
    )
    sobrantes = sorted(set(_PATRON_CODIGO_GENERICO.findall(texto_completo)))

    return buffer.getvalue(), sobrantes
