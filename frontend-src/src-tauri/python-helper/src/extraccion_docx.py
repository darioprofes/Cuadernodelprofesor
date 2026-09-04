# ==========================================================
# Extracción de .docx a Markdown (Herramientas IA)
# ==========================================================
#
# Convierte un documento Word a texto Markdown en vez de solo volcar texto
# plano, para que las tablas (p.ej. una tabla de calificaciones) lleguen a la
# IA online con su estructura de filas/columnas intacta -- python-docx
# expone párrafos y tablas como dos listas separadas
# (doc.paragraphs / doc.tables), perdiendo el orden original del documento,
# así que se recorre el XML del cuerpo directamente para intercalarlos tal
# como aparecen. Puramente extracción: no persiste nada, no llama a ninguna
# IA (mismo criterio que services/horario_pdf.py con el PDF del horario).

import io

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

# Los nombres de estilo varían según el idioma de la plantilla de Word
# ("Heading 1" en plantillas en inglés, "Título 1" en plantillas en
# español) -- se contemplan ambos.
_NIVEL_TITULO = {
    "Title": 1, "Título": 1,
    "Heading 1": 1, "Título 1": 1,
    "Heading 2": 2, "Título 2": 2,
    "Heading 3": 3, "Título 3": 3,
    "Heading 4": 4, "Título 4": 4,
}


def _iter_bloques(documento):

    # doc.element.body.iterchildren() da los hijos en el orden real del
    # documento; se envuelve cada uno en el objeto de alto nivel de
    # python-docx correspondiente según su etiqueta XML.
    for hijo in documento.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            yield Paragraph(hijo, documento)
        elif hijo.tag == qn("w:tbl"):
            yield Table(hijo, documento)


def _escapar_asteriscos(texto):

    # Un "*" suelto en el texto original (p.ej. una llamada a nota al pie)
    # rompería el markdown que se genera envolviendo negrita/cursiva con
    # ** y * -- se escapa antes de aplicar ese envoltorio.
    return texto.replace("*", "\\*")


def _parrafo_a_markdown(paragraph):

    partes = []

    for run in paragraph.runs:
        texto = _escapar_asteriscos(run.text)

        if not texto:
            continue

        if run.bold and run.italic:
            texto = f"***{texto}***"
        elif run.bold:
            texto = f"**{texto}**"
        elif run.italic:
            texto = f"*{texto}*"

        partes.append(texto)

    return "".join(partes)


def _celda_a_texto(cell):

    # Una celda puede tener varios párrafos internos; se unen con un
    # espacio porque una tabla Markdown es de una sola línea por fila.
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _tabla_a_markdown(table):

    filas_md = []

    for i, row in enumerate(table.rows):
        celdas = [_celda_a_texto(cell).replace("|", "\\|") for cell in row.cells]
        filas_md.append("| " + " | ".join(celdas) + " |")

        if i == 0:
            filas_md.append("| " + " | ".join(["---"] * len(celdas)) + " |")

    return "\n".join(filas_md)


def extraer_markdown_docx(contenido_bytes):

    documento = Document(io.BytesIO(contenido_bytes))
    bloques = []

    for bloque in _iter_bloques(documento):

        if isinstance(bloque, Table):
            bloques.append(_tabla_a_markdown(bloque))
            continue

        texto = _parrafo_a_markdown(bloque)

        if not texto.strip():
            continue

        nivel = _NIVEL_TITULO.get(bloque.style.name) if bloque.style else None

        if nivel:
            texto = "#" * nivel + " " + texto

        bloques.append(texto)

    return "\n\n".join(bloques)
