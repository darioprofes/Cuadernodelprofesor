# ==========================================================
# Anonimizador de documentos (Herramientas IA)
# ==========================================================
#
# Detecta datos personales en un texto pegado a mano (actas de evaluación,
# informes...) y los sustituye por códigos aleatorios, para poder pegar el
# documento en una IA online (Claude, ChatGPT) sin enviarle datos reales.
# No persiste nada: el mapa código -> dato real se devuelve una única vez en
# la respuesta y es responsabilidad del frontend guardarlo solo en memoria
# del navegador (nunca en Postgres ni en localStorage) para poder reintegrar
# la respuesta de la IA online más tarde.
#
# spaCy (NER) es la herramienta correcta para "encontrar entidades borrosas"
# (personas, organizaciones, lugares) -- un LLM de chat no aporta nada aquí
# y sería mucho más lento. Los patrones estructurados y predecibles (DNI,
# direcciones, nivel+grupo, cargos únicos de centro) se detectan aparte con
# regex, porque el NER los reconoce mal.

import json
import re
import secrets
from pathlib import Path

import spacy

_BASE = Path(__file__).parent
_nlp = spacy.load("es_core_news_md")

_MATERIAS_CONOCIDAS = {
    m.lower() for m in json.loads((_BASE / "materias_oficiales.json").read_text(encoding="utf-8"))
}

# Términos NEAE/diagnósticos habituales (dislexia, TDAH...): el NER los
# confunde a veces con nombres propios por ir capitalizados y a menudo
# pegados a paréntesis/dos puntos en textos de atención educativa (visto en
# real: "Necesidades NEAE: Dislexia" se comía "Dislexia" como si fuera un
# nombre). Mismo mecanismo que _MATERIAS_CONOCIDAS -- lista curada a mano,
# no exhaustiva; cualquier término que falte aquí sigue cubierto por la
# revisión manual del profesor antes de mandar nada fuera.
_TERMINOS_NEAE_CONOCIDOS = {
    t.lower() for t in json.loads((_BASE / "neae_terminos.json").read_text(encoding="utf-8"))
}

_PATRON_DNI = re.compile(r"\b\d{8}[A-Za-z]\b")
_PATRON_DIRECCION = re.compile(
    r"\b(?:[Cc]alle|[Aa]venida|[Aa]vda\.?|[Pp]laza|[Pp]aseo|[Cc]/)\s+[^\n,]+?\s+\d+"
    r"(?:,\s*\d+º?[A-Za-z]?)?"
)
_PATRON_CODIGO_POSTAL = re.compile(r"\b\d{5}\b")
# Acotado a la MISMA LÍNEA (nada de \n, coma ni punto) y con longitud máxima:
# un \s+ sin acotar cruza saltos de línea y se "come" varias líneas del
# documento por error (visto en pruebas reales).
_PATRON_CENTRO = re.compile(r"\b(?:IES|CEIP|CRA|CPI|EOI|CIFP|Colegio|Instituto)\b[^\n,\.]{0,30}")
# Nivel + grupo (p.ej. "2º ESO B"): identificador indirecto -- combinado con
# fecha/materia puede acotar a muy poca gente. Prefijo GRUPO_ en vez de
# PERS_ porque no es una persona ni una organización.
_PATRON_NIVEL = re.compile(
    r"\b[1-4]º\s*(?:de\s+)?ESO(?:\s+[A-ZÑ](?!\w))?"
    r"|\b[1-2]º\s*(?:de\s+)?Bachillerato(?:\s+[A-ZÑ](?!\w))?"
)
# En un centro normalmente solo hay una persona con cada uno de estos
# cargos: combinado con el nombre del centro, identifica tanto como un
# nombre propio. Lista curada a mano porque es un conjunto pequeño y
# conocido, no algo que un NER detecte.
_CARGOS_UNICOS = [
    "jefe de estudios", "jefa de estudios", "director", "directora",
    "secretario", "secretaria", "orientador", "orientadora",
]

INSTRUCCION_IA_ONLINE = (
    "IMPORTANTE: este documento contiene códigos con el formato PERS_XXXXXX y "
    "GRUPO_XXXXXX (por ejemplo: PERS_A63425, GRUPO_77DF7E). Son sustitutos de "
    "datos personales reales, generados para proteger la privacidad. NO los "
    "traduzcas, no los modifiques, no los interpretes ni les des otro "
    "significado -- consérvalos EXACTAMENTE tal cual aparecen, letra por "
    "letra, en tu respuesta. Los necesito intactos para reconstruir el "
    "documento con los datos reales después."
)


def _detectar_candidatos(texto):

    doc = _nlp(texto)

    candidatos = [(ent.start_char, ent.end_char, ent.text)
                  for ent in doc.ents if ent.label_ in ("PER", "ORG", "LOC")]

    for patron in (_PATRON_DNI, _PATRON_DIRECCION, _PATRON_CODIGO_POSTAL):
        for m in patron.finditer(texto):
            candidatos.append((m.start(), m.end(), m.group()))

    for m in _PATRON_CENTRO.finditer(texto):
        # rstrip() para no arrastrar espacios sobrantes -- recalculando
        # "fin" a partir del texto ya recortado, para que el rango siga
        # coincidiendo exactamente con lo que se sustituirá después.
        recortado = m.group().rstrip()
        candidatos.append((m.start(), m.start() + len(recortado), recortado))

    niveles_detectados = [(m.start(), m.end(), m.group()) for m in _PATRON_NIVEL.finditer(texto)]
    candidatos += niveles_detectados

    for cargo in _CARGOS_UNICOS:
        for m in re.finditer(r"\b" + re.escape(cargo) + r"\b", texto, re.IGNORECASE):
            candidatos.append((m.start(), m.end(), m.group()))

    return candidatos, niveles_detectados


def _recortar_saltos_de_linea(candidatos):

    # Ningún candidato legítimo debería cruzar un salto de línea en un
    # documento de este tipo -- tanto el NER como algún regex pueden
    # producir spans que fusionan líneas distintas (visto en pruebas
    # reales con encabezados densos). Se recorta al primer salto de línea
    # en vez de descartar entero, para no perder la parte válida.
    recortados = []

    for inicio, fin, texto in candidatos:
        if "\n" in texto:
            texto = texto.split("\n")[0].rstrip()
            fin = inicio + len(texto)
            if not texto.strip():
                continue
        recortados.append((inicio, fin, texto))

    return recortados


def _resolver_solapamientos(candidatos):

    # P.ej. el patrón de código postal (5 dígitos) puede encontrar una
    # coincidencia DENTRO de un DNI (8 dígitos + letra). Se procesa por
    # orden de inicio y, entre solapados, gana el rango más largo (más
    # específico) -- sin esto, sustituir rangos que se pisan corrompe el
    # texto.
    ordenados = sorted(candidatos, key=lambda c: (c[0], -(c[1] - c[0])))

    resueltos = []
    fin_ocupado_hasta = -1

    for inicio, fin, texto in ordenados:
        if inicio < fin_ocupado_hasta:
            continue
        resueltos.append((inicio, fin, texto))
        fin_ocupado_hasta = fin

    return resueltos


def _generar_codigos(texto):
    """Detecta datos personales en `texto` y genera un código aleatorio por
    cada texto único detectado (no por aparición). Devuelve
    (candidatos, codigo_de, mapa_real) -- `candidatos` conserva las
    posiciones (para sustituir por offset en texto plano), `codigo_de`/
    `mapa_real` son el mismo mapeo en las dos direcciones (para sustituir
    por texto literal, como necesita la variante .docx más abajo)."""

    candidatos, niveles_detectados = _detectar_candidatos(texto)
    candidatos = _recortar_saltos_de_linea(candidatos)

    # El recorte por salto de línea a veces deja suelta una sola palabra
    # que en realidad era ruido de una entidad LOC/ORG mal delimitada por
    # spaCy (p.ej. "Física y Química" fusionada con la línea siguiente,
    # dejando "Química" sola tras el recorte sin ser un dato personal).
    # Subir el umbral de confianza del NER no es viable aquí: es_core_news_md
    # no expone una puntuación de confianza fiable por entidad sin
    # reconfigurar el pipeline entero. Una lista de exclusión con las
    # materias reales del propio cuaderno es más simple y directa.
    candidatos = [
        (i, f, t) for i, f, t in candidatos
        if t.strip().lower() not in _MATERIAS_CONOCIDAS
        and t.strip().lower() not in _TERMINOS_NEAE_CONOCIDOS
    ]

    candidatos = _resolver_solapamientos(candidatos)

    textos_nivel = {texto_n for _, _, texto_n in niveles_detectados}

    codigo_de = {}
    mapa_real = {}
    for _, _, texto_candidato in candidatos:
        if texto_candidato not in codigo_de:
            prefijo = "GRUPO" if texto_candidato in textos_nivel else "PERS"
            codigo = "%s_%s" % (prefijo, secrets.token_hex(3).upper())
            codigo_de[texto_candidato] = codigo
            mapa_real[codigo] = texto_candidato

    return candidatos, codigo_de, mapa_real


def anonimizar(texto):
    """Devuelve (documento_para_pegar, mapa) donde `mapa` es código -> dato
    real. `documento_para_pegar` incluye ya la instrucción para la IA online
    delante del texto anonimizado."""

    candidatos, codigo_de, mapa_real = _generar_codigos(texto)

    anonimizado = texto
    for inicio, fin, texto_candidato in sorted(candidatos, key=lambda c: -c[0]):
        anonimizado = anonimizado[:inicio] + codigo_de[texto_candidato] + anonimizado[fin:]

    # Pasada de repaso: por si algún dato detectado aparece OTRA VEZ en el
    # texto sin que el NER lo haya etiquetado esa vez concreta (limitación
    # típica de NER frente a coincidencia exacta).
    for texto_candidato, codigo in sorted(codigo_de.items(), key=lambda x: -len(x[0])):
        anonimizado = anonimizado.replace(texto_candidato, codigo)

    documento_para_pegar = INSTRUCCION_IA_ONLINE + "\n\n" + anonimizado

    return documento_para_pegar, mapa_real


_PATRON_CODIGO_GENERICO = re.compile(r"\b(?:PERS|GRUPO)_[0-9A-F]{6}\b")


def _sustituir_codigos_en_parrafo(paragraph, patron, mapa):

    # Se sustituye run por run, no en el texto completo del párrafo, para no
    # perder el formato (negrita, cursiva...) que la IA online le haya dado
    # a lo que rodea el código: cambiar run.text conserva su estilo, un
    # find-replace sobre el texto completo del párrafo lo destruiría. Un
    # código partido entre dos runs (p.ej. por un cambio de estilo a mitad)
    # no se detecta aquí -- se queda sin resolver y se avisa en el resultado,
    # en vez de intentar fusionar runs con el riesgo de romper el documento.
    for run in paragraph.runs:
        if run.text:
            run.text = patron.sub(lambda m: mapa[m.group(0)], run.text)


def _sustituir_textos_en_parrafo(paragraph, pares_texto_codigo):

    # Mismo criterio que _sustituir_codigos_en_parrafo: run por run, no en
    # el texto completo del párrafo, para conservar negrita/cursiva/tablas
    # tal cual estaban. Un dato detectado que quede partido entre dos runs
    # (p.ej. un nombre con un cambio de estilo a mitad) no se anonimizaría
    # -- limitación aceptada, igual que en reintegrar_docx.
    for run in paragraph.runs:
        if not run.text:
            continue
        for texto_candidato, codigo in pares_texto_codigo:
            if texto_candidato in run.text:
                run.text = run.text.replace(texto_candidato, codigo)


def anonimizar_docx(contenido_bytes):
    """Recibe un .docx con datos personales (acta, informe...) y devuelve
    (bytes_docx, mapa) con los datos sustituidos por códigos SIN tocar el
    formato original -- útil para poder pegar/adjuntar el propio .docx en
    una IA online (p.ej. Claude generando su respuesta también en .docx)
    en vez de convertirlo antes a texto plano."""

    import io

    from docx import Document

    documento = Document(io.BytesIO(contenido_bytes))

    bloques_texto = [p.text for p in documento.paragraphs]
    bloques_texto += [
        celda.text
        for tabla in documento.tables
        for fila in tabla.rows
        for celda in fila.cells
    ]
    texto_completo = "\n".join(bloques_texto)

    _, codigo_de, mapa_real = _generar_codigos(texto_completo)

    if codigo_de:
        # Más largo primero: si un texto detectado es substring de otro
        # más largo (p.ej. un nombre de pila que coincide con el inicio
        # del nombre completo), sustituir el más largo antes evita dejarlo
        # a medio anonimizar.
        pares = sorted(codigo_de.items(), key=lambda x: -len(x[0]))

        for parrafo in documento.paragraphs:
            _sustituir_textos_en_parrafo(parrafo, pares)

        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        _sustituir_textos_en_parrafo(parrafo, pares)

    if documento.paragraphs:
        documento.paragraphs[0].insert_paragraph_before(INSTRUCCION_IA_ONLINE)
    else:
        documento.add_paragraph(INSTRUCCION_IA_ONLINE)

    buffer = io.BytesIO()
    documento.save(buffer)

    return buffer.getvalue(), mapa_real


def reintegrar_docx(contenido_bytes, mapa):
    """Recibe el .docx que ha devuelto la IA online (con los códigos
    PERS_/GRUPO_ intactos) y un .docx con los mismos datos reales, sin
    tocar el formato que le haya dado la IA. Devuelve (bytes_docx, sobrantes)
    donde `sobrantes` son códigos que no se han podido resolver (normalmente
    porque quedaron partidos entre dos runs de estilo distinto)."""

    import io

    from docx import Document

    documento = Document(io.BytesIO(contenido_bytes))

    if mapa:
        patron = re.compile("|".join(re.escape(codigo) for codigo in mapa))

        for parrafo in documento.paragraphs:
            _sustituir_codigos_en_parrafo(parrafo, patron, mapa)

        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        _sustituir_codigos_en_parrafo(parrafo, patron, mapa)

    buffer = io.BytesIO()
    documento.save(buffer)

    texto_completo = "\n".join(p.text for p in documento.paragraphs)
    texto_completo += "\n".join(
        celda.text
        for tabla in documento.tables
        for fila in tabla.rows
        for celda in fila.cells
    )
    sobrantes = sorted(set(_PATRON_CODIGO_GENERICO.findall(texto_completo)))

    return buffer.getvalue(), sobrantes
